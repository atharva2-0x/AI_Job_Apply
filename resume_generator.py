# resume_generator.py
import openai
import json
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from config import OPENAI_API_KEY
import pdfkit  # For PDF conversion

openai.api_key = OPENAI_API_KEY

class ResumeGenerator:
    def __init__(self, user_data_path="user_data.json", resume_template="templates/resume_template.docx"):
        """
        Initialize the ResumeGenerator with user data and resume template.
        """
        with open(user_data_path) as f:
            self.user_data = json.load(f)
        self.resume_template = resume_template
    
    def analyze_job_description(self, job_description):
        """
        Analyze the job description and extract key requirements.
        """
        prompt = f"""
        Analyze this job description and extract:
        - Required technical/hard skills (e.g., Python, AWS)
        - Soft skills (e.g., Team Leadership)
        - Tools/frameworks (e.g., React, TensorFlow)
        - Key responsibilities
        Job Description: {job_description}
        Format output as JSON: {{"skills": [], "tools": [], "responsibilities": []}}
        """
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        return json.loads(response.choices[0].message.content)
    
    def update_resume(self, job_analysis, output_format="pdf"):
        """
        Update the user-provided resume based on job analysis.
        """
        # Load the resume template
        doc = Document(self.resume_template)
        
        # Update sections
        self._update_summary(doc, job_analysis)
        self._update_skills(doc, job_analysis)
        self._update_experience(doc, job_analysis)
        self._update_projects(doc, job_analysis)
        
        # Save the updated resume
        output_file = "tailored_resume.docx"
        doc.save(output_file)
        
        # Convert to PDF if required
        if output_format == "pdf":
            self._convert_to_pdf(output_file, "tailored_resume.pdf")
            return "tailored_resume.pdf"
        return output_file
    
    def _update_summary(self, doc, job_analysis):
        """
        Update the summary section.
        """
        prompt = f"""
        Write a 3-line professional summary incorporating:
        - The candidate's background: {self.user_data['professional']['summary']}
        - Job requirements: {job_analysis}
        - Use active verbs and quantifiable achievements.
        """
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        summary = response.choices[0].message.content
        
        # Replace the summary section in the document
        for paragraph in doc.paragraphs:
            if "SUMMARY" in paragraph.text:
                paragraph.text = f"SUMMARY\n{summary}"
                break
    
    def _update_skills(self, doc, job_analysis):
        """
        Update the skills section.
        """
        user_skills = self.user_data['professional']['skills']
        required_skills = job_analysis['skills']
        
        # Prioritize skills that match job requirements
        tailored_skills = [skill for skill in user_skills if skill in required_skills]
        
        # Add missing skills if necessary
        for skill in required_skills:
            if skill not in tailored_skills:
                tailored_skills.append(skill)
        
        # Replace the skills section in the document
        for paragraph in doc.paragraphs:
            if "SKILLS" in paragraph.text:
                paragraph.text = f"SKILLS\n{', '.join(tailored_skills)}"
                break
    
    def _update_experience(self, doc, job_analysis):
        """
        Update the experience section.
        """
        tailored_experience = []
        for exp in self.user_data['professional']['experience']:
            prompt = f"""
            Rewrite this work experience to align with the job responsibilities:
            - Job Responsibilities: {job_analysis['responsibilities']}
            - Original Experience: {exp}
            - Focus on relevant achievements and skills.
            Format output as JSON: {{"title": "", "company": "", "duration": "", "bullets": []}}
            """
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}]
            )
            tailored_experience.append(json.loads(response.choices[0].message.content))
        
        # Replace the experience section in the document
        for paragraph in doc.paragraphs:
            if "WORK EXPERIENCE" in paragraph.text:
                experience_text = "WORK EXPERIENCE\n"
                for exp in tailored_experience:
                    experience_text += f"{exp['title']} at {exp['company']} ({exp['duration']})\n"
                    for bullet in exp['bullets']:
                        experience_text += f"- {bullet}\n"
                paragraph.text = experience_text
                break
    
    def _update_projects(self, doc, job_analysis):
        """
        Update the projects section.
        """
        tailored_projects = []
        for project in self.user_data['professional'].get('projects', []):
            prompt = f"""
            Rewrite this project to align with the job requirements:
            - Job Requirements: {job_analysis}
            - Original Project: {project}
            - Focus on relevant tools and outcomes.
            Format output as JSON: {{"title": "", "description": "", "tools": []}}
            """
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[{"role": "user", "content": prompt}]
            )
            tailored_projects.append(json.loads(response.choices[0].message.content))
        
        # Replace the projects section in the document
        for paragraph in doc.paragraphs:
            if "PROJECTS" in paragraph.text:
                projects_text = "PROJECTS\n"
                for project in tailored_projects:
                    projects_text += f"{project['title']}\n"
                    projects_text += f"- {project['description']}\n"
                    projects_text += f"- Tools: {', '.join(project['tools'])}\n"
                paragraph.text = projects_text
                break
    
    def _convert_to_pdf(self, docx_file, pdf_file):
        """
        Convert a Word document to PDF.
        """
        pdfkit.from_file(docx_file, pdf_file)
